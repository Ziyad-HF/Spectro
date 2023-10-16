from PyQt5.QtWidgets import QMainWindow, QInputDialog  # QApplication
from PyQt5.uic import loadUiType
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from classes import *
import sys
from os import listdir, path, remove
from docx import Document
from docx.shared import Inches
from PIL import Image
from docx2pdf import convert

FORM_CLASS, _ = loadUiType(path.join(path.dirname(__file__), "mainWindow.ui"))


def add_snapshots_to_report(paths, snapshots_signals, doc, number):
    if len(paths) != 0:
        # for loop to add graph two images
        for imgPath, snapshotSignals in zip(paths, snapshots_signals):
            doc.add_paragraph("\n")
            doc.add_heading(f"Snapshot {number}", 3)
            number += 1
            img = Image.open(imgPath)
            width, height = img.size
            doc.paragraphs[-1].add_run().add_picture(imgPath, width=Inches(7),
                                                     height=Inches(height/100))
            doc.add_paragraph("Snapshot Statistics")
            num_rows = 1 + len(snapshotSignals.keys())
            num_cols = 6
            table = doc.add_table(rows=num_rows, cols=num_cols)
            table.style = 'Light Shading'
            headings = ['Signal', 'Max', 'Min', 'Mean', 'STD', 'Duration']
            for i, heading in enumerate(headings):
                cell = table.cell(0, i)
                cell.text = heading
                cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = cell.paragraphs[0].runs[0]
                run.bold = True
                run.font.size = Pt(12)

            row_index = 1
            for signal in snapshotSignals.items():
                signal_title, signal_stats = signal
                data = [signal_title]
                data.extend(signal_stats)
                for col, text in enumerate(data):
                    cell = table.cell(row_index, col)
                    cell.text = str(text)
                    cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    cell.width = Inches(2.0)
                    cell.height = Inches(1.0)
                row_index += 1
    # QApplication.processEvents()


class MainApp(QMainWindow, FORM_CLASS):

    def __init__(self, parent=None):
        super(MainApp, self).__init__(parent)
        QMainWindow.__init__(self, parent=None)
        self.setupUi(self)
        self.setWindowTitle("Spectro")
        self.setWindowIcon(QtGui.QIcon('icons/logo.ico'))
        self.graphicsView.setLimits(xMin=0)
        self.graphicsView2.setLimits(xMin=0)
        self.graphicsView.setMenuEnabled(False)
        self.graphicsView2.setMenuEnabled(False)
        self.moveDownBtn.setIcon(QtGui.QIcon(QtGui.QPixmap("icons/move down.png")))
        self.moveUpBtn.setIcon(QtGui.QIcon(QtGui.QPixmap("icons/move up.png")))
        self.linkUnlinkBtn.setIcon(QtGui.QIcon(QtGui.QPixmap("icons/link.png")))
        self.linkStatus = False  # False for unlinked True for linked
        # creating the first graph
        self.graphOne = Graph(self.graphicsView, self, [self.addSignalBtn, self.pausePlayBtn,
                                                        self.zoomInBtn, self.zoomOutBtn,
                                                        self.resetBtn, self.clearBtn,
                                                        self.snapshotBtn, self.showHideBtn,
                                                        self.addTitleBtn, self.deleteSignalBtn,
                                                        self.speedUpBtn, self.speedDownBtn,
                                                        self.syncUnsyncSignalsBtn,
                                                        self.signalComboBox, self.colorBtn,
                                                        self.addTitleLineEdit, self.originalViewBtn,
                                                        self.defaultSpeedBtn, self.actionGenerate_PDF,
                                                        self.moveDownBtn
                                                        ], 1)
        # creating the second graph
        self.graphTwo = Graph(self.graphicsView2, self, [self.addSignalBtn2, self.pausePlayBtn2,
                                                         self.zoomInBtn2, self.zoomOutBtn2,
                                                         self.resetBtn2, self.clearBtn2,
                                                         self.snapshotBtn2, self.showHideBtn2,
                                                         self.addTitleBtn2, self.deleteSignalBtn2,
                                                         self.speedUpBtn2, self.speedDownBtn2,
                                                         self.syncUnsyncSignalsBtn2,
                                                         self.signalComboBox2, self.colorBtn2,
                                                         self.addTitleLineEdit2, self.originalViewBtn2,
                                                         self.defaultSpeedBtn2, self.actionGenerate_PDF,
                                                         self.moveUpBtn
                                                         ], 2)
        self.handle_buttons()
        self.linkTimer = QtCore.QTimer()
        self.linkGraphSpeed = 21

    def handle_buttons(self):
        self.actionGenerate_PDF.triggered.connect(self.export_to_pdf)
        self.linkUnlinkBtn.clicked.connect(self.link_graphs)
        self.moveUpBtn.clicked.connect(self.move_signal_up)
        self.moveDownBtn.clicked.connect(self.move_signal_down)
        self.actionGenerate_PDF.setEnabled(False)

    def link_graphs(self):
        if self.linkStatus:
            self.linkUnlinkBtn.setIcon(QtGui.QIcon(QtGui.QPixmap("icons/link.png")))
            self.linkUnlinkBtn.setText("link")
            self.linkTimer.stop()
            self.graphicsView2.setXLink(self.graphicsView2)
            self.graphicsView2.setYLink(self.graphicsView2)
            self.addSignalBtn2.disconnect()
            self.addSignalBtn.disconnect()
            self.pausePlayBtn.disconnect()
            self.zoomInBtn.disconnect()
            self.zoomOutBtn.disconnect()
            self.resetBtn.disconnect()
            self.speedUpBtn.disconnect()
            self.speedDownBtn.disconnect()
            self.defaultSpeedBtn.disconnect()
            self.graphOne.link(False)
            self.graphTwo.link(False)
            self.pausePlayBtn2.show()
            self.zoomInBtn2.show()
            self.zoomOutBtn2.show()
            self.resetBtn2.show()
            self.speedUpBtn2.show()
            self.speedDownBtn2.show()
            self.defaultSpeedBtn2.show()
            self.syncUnsyncSignalsBtn.show()
            self.syncUnsyncSignalsBtn2.show()
            self.linkStatus = False
            self.graphOne.timer.start()
            self.graphTwo.timer.start()

        else:
            if self.graphOne.timer.isActive():
                self.graphOne.timer.stop()
            if self.graphTwo.timer.isActive():
                self.graphTwo.timer.stop()
            self.linkUnlinkBtn.setIcon(QtGui.QIcon(QtGui.QPixmap("icons/unlink.png")))
            self.linkUnlinkBtn.setText("Unlink")
            self.graphOne.link(True)
            self.graphTwo.link(True)
            self.graphTwo.speedOfGraph = self.graphOne.speedOfGraph
            self.linkGraphSpeed = self.graphOne.speedOfGraph
            self.linkTimer.setInterval(self.linkGraphSpeed)
            self.linkTimer.timeout.connect(self.link_plot)
            self.graphTwo.plottingPoint = self.graphOne.plottingPoint
            # hide buttons
            self.pausePlayBtn2.hide()
            self.zoomInBtn2.hide()
            self.zoomOutBtn2.hide()
            self.resetBtn2.hide()
            self.speedUpBtn2.hide()
            self.speedDownBtn2.hide()
            self.defaultSpeedBtn2.hide()
            self.syncUnsyncSignalsBtn.hide()
            self.syncUnsyncSignalsBtn2.hide()
            self.addSignalBtn2.clicked.connect(self.link_add_signal_graph2)
            self.addSignalBtn.clicked.connect(self.link_add_signal_graph1)
            self.pausePlayBtn.clicked.connect(self.link_pause_play_graph)
            self.zoomInBtn.clicked.connect(self.link_zoom_in)
            self.zoomOutBtn.clicked.connect(self.link_zoom_out)
            self.resetBtn.clicked.connect(self.link_reset_graph)
            self.speedUpBtn.clicked.connect(self.link_graph_speed_up)
            self.speedDownBtn.clicked.connect(self.link_graph_speed_down)
            self.defaultSpeedBtn.clicked.connect(self.link_default_speed)
            self.graphicsView2.setXLink(self.graphicsView)
            self.graphicsView2.setYLink(self.graphicsView)
            self.linkStatus = True
            self.linkTimer.start()

    def link_pause_play_graph(self):
        if self.linkTimer.isActive():
            self.pausePlayBtn.setText("Play")
            self.pausePlayBtn.setIcon(QtGui.QIcon(QtGui.QPixmap("icons/play.png")))
            self.linkTimer.stop()
        else:
            self.pausePlayBtn.setText("Pause")
            self.pausePlayBtn.setIcon(QtGui.QIcon(QtGui.QPixmap("icons/pause.png")))
            self.linkTimer.start()


    def link_add_signal_graph1(self):
        self.linkTimer.stop()
        self.graphOne.add_signal()
        self.linkTimer.start()
        # QApplication.processEvents()

    def link_add_signal_graph2(self):
        self.linkTimer.stop()
        self.graphTwo.add_signal()
        self.linkTimer.start()

    def link_zoom_in(self):
        self.graphOne.zoom_in()
        self.graphTwo.zoom_in()

    def link_zoom_out(self):
        self.graphOne.zoom_out()
        self.graphTwo.zoom_out()

    def link_reset_graph(self):
        self.graphOne.reset_graph()
        self.graphTwo.reset_graph()
        if not self.linkTimer.isActive():
            self.linkTimer.start()

    def link_graph_speed_up(self):
        self.linkGraphSpeed -= 5
        self.linkTimer.setInterval(self.linkGraphSpeed)
        self.check_speed_in_range()

    def link_graph_speed_down(self):
        self.linkGraphSpeed += 5
        self.linkTimer.setInterval(self.linkGraphSpeed)
        self.check_speed_in_range()

    def link_default_speed(self):
        self.linkGraphSpeed = 21
        self.linkTimer.setInterval(self.linkGraphSpeed)
        self.check_speed_in_range()

    def check_speed_in_range(self):
        if self.linkGraphSpeed > 5:
            self.speedUpBtn.setEnabled(True)
        else:
            self.speedUpBtn.setEnabled(False)

        if self.linkGraphSpeed < 100:
            self.speedDownBtn.setEnabled(True)
        else:
            self.speedDownBtn.setEnabled(False)

    def link_plot(self):
        self.graphOne.plot()
        self.graphTwo.plot()

    def clean(self):
        graph_one_folder_path = "temp/imgs/graphOne"
        graph_two_folder_path = "temp/imgs/graphTwo"
        # List all files in the folder
        graph_one_files = listdir(graph_one_folder_path)
        graph_two_files = listdir(graph_two_folder_path)

        for file in graph_one_files:
            file_path = path.join(graph_one_folder_path, file)
            # Check if the path is a file (not a subdirectory)
            if path.isfile(file_path):
                remove(file_path)
        for file in graph_two_files:
            file_path = path.join(graph_two_folder_path, file)
            # Check if the path is a file (not a subdirectory)
            if path.isfile(file_path):
                remove(file_path)  # Delete the file

        self.graphOne.snapshotsPaths.clear()
        self.graphOne.imageNumber = 0
        self.graphTwo.snapshotsPaths.clear()
        self.graphTwo.imageNumber = 0

    def create_word(self):
        doc = Document()
        # remove doc padding and margins
        doc.sections[0].left_margin = Inches(0.8)
        doc.sections[0].right_margin = Inches(0.8)
        doc.sections[0].top_margin = Inches(0.8)
        signal_dict_graph_one = self.graphOne.signalDictionary
        signal_dict_graph_two = self.graphTwo.signalDictionary
        graph_one_signals = list(signal_dict_graph_one.values())
        graph_two_signals = list(signal_dict_graph_two.values())
        doc.add_heading("Report of statistics and Snapshots", 0)
        doc.add_paragraph("")

        doc.add_heading('General statistics on signals', 1)
        doc.add_paragraph("")

        # Define the number of rows and columns for the table
        num_rows = len(graph_one_signals) + len(graph_two_signals) + 1  # number of signal
        num_cols = 6
        # Add a table with the specified number of rows and columns
        table = doc.add_table(rows=num_rows, cols=num_cols)
        # Set the style of the table
        table.style = 'Light Shading'
        # Define table headings
        headings = ['Signal', 'Max', 'Min', 'Mean', 'STD', 'Duration']

        # Populate the first row of the table with headings
        for i, heading in enumerate(headings):
            cell = table.cell(0, i)
            cell.text = heading
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = cell.paragraphs[0].runs[0]
            run.bold = True
            run.font.size = Pt(12)

        # Add signals stats numbers
        data = []
        for signal in graph_one_signals:
            data.append([signal.title] + signal.stats)

        for signal in graph_two_signals:
            data.append([signal.title] + signal.stats)

        for i, row_data in enumerate(data):
            for j, text in enumerate(row_data):
                cell = table.cell(i + 1, j)
                cell.text = str(text)
                cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                cell.width = Inches(2.0)
                cell.height = Inches(1.0)

        doc.add_paragraph("")
        if len(self.graphOne.snapshotsPaths) or len(self.graphTwo.snapshotsPaths) != 0:
            doc.add_heading("Here are Each snapshot and its statistics:\n")
        else:``
            doc.add_heading("No snapshots were taken", 1)
        add_snapshots_to_report(self.graphOne.snapshotsPaths, self.graphOne.snapshotStats, doc, 1)
        add_snapshots_to_report(self.graphTwo.snapshotsPaths, self.graphTwo.snapshotStats, doc,
                                len(self.graphOne.snapshotsPaths) + 1)

        save_path = QFileDialog.getSaveFileName(self, 'Save File', "Report", "PDf Files (*.pdf)")[0]
        print(save_path)
        doc.save(f"{save_path}.docx")
        return f"{save_path}.docx"

    def export_to_pdf(self):
        save_path = self.create_word()
        # Converting word files to PDFs
        convert(save_path)
        remove(save_path)
        self.clean()
        QMessageBox.information(self, "Saved", "Report saved successfully")

    def move_signal_up(self):
        file_path, title,color = self.graphTwo.signalDictionary[self.signalComboBox2.currentText()].return_path_title()
        self.graphOne.add_signal(file_path, title,color)
        self.graphTwo.delete_signal(True)

    def move_signal_down(self):
        file_path, title,color = self.graphOne.signalDictionary[self.signalComboBox.currentText()].return_path_title()
        self.graphTwo.add_signal(file_path, title,color)
        self.graphOne.delete_signal(True)


def main():
    app = QApplication(sys.argv)
    window = MainApp()
    window.show()
    app.exec_()


if __name__ == '__main__':
    main()
